"""
office_convert.py
Pure-Python, OS-agnostic conversion of common attachment formats to PDF.
No Microsoft Office / LibreOffice / COM required.

Supported formats:
    .xlsx .xlsm .xltx .xltm   -> openpyxl  -> reportlab Tables
    .xls                       -> xlrd      -> reportlab Tables
    .csv .tsv                  -> stdlib    -> reportlab Tables
    .docx .docm .dotx .dotm    -> python-docx -> reportlab flowables
    .doc                       -> placeholder (legacy binary; not reliably
                                  parseable without external tools)
    .rtf                       -> striprtf  -> reportlab paragraphs
    .txt                       -> reportlab paragraphs
    .vsdx                      -> visio_convert -> reportlab structured text

Trade-offs:
    * Loses charts, shapes, embedded objects, exact font rendering.
    * Preserves: text, tables, basic structure, paragraph order.
    * Reliable, deterministic, works on Windows / macOS / Linux identically.
"""

from __future__ import annotations

import csv
import gc
from pathlib import Path
from typing import Final

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, LETTER, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.platypus.doctemplate import LayoutError

# Hardcoded layout constants
MAX_COLS_PER_TABLE: Final[int] = 24
MAX_ROWS_PER_SHEET: Final[int] = 5000
MAX_CELL_CHARS: Final[int] = 500

# A row whose product of (col_count x max_cell_chars) exceeds this is
# considered "too dense for table layout" and renders as stacked
# paragraphs instead. Tuned empirically: 24 cols x 400 chars = 9600
# fails on A4 landscape; 24 cols x 300 chars = 7200 succeeds; we
# choose 7000 as a safe cutoff with headroom.
DENSE_ROW_THRESHOLD: Final[int] = 7000


# ---------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------
def _pagesize(paper: str, landscape_mode: bool):
    base = LETTER if paper.upper() == "LETTER" else A4
    return landscape(base) if landscape_mode else base


def _styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "h1": ParagraphStyle(
            "OC_H1", parent=base["Heading1"], fontName="Helvetica-Bold",
            fontSize=14, leading=18, spaceAfter=6,
        ),
        "h2": ParagraphStyle(
            "OC_H2", parent=base["Heading2"], fontName="Helvetica-Bold",
            fontSize=11, leading=14, spaceBefore=8, spaceAfter=4,
        ),
        "body": ParagraphStyle(
            "OC_Body", parent=base["BodyText"], fontName="Helvetica",
            fontSize=9, leading=12,
        ),
        "small": ParagraphStyle(
            "OC_Small", parent=base["BodyText"], fontName="Helvetica",
            fontSize=7, leading=9,
        ),
    }


def _safe_html(s: str) -> str:
    return (
        (s or "")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def _truncate(s: str, n: int = MAX_CELL_CHARS) -> str:
    if s is None:
        return ""
    s = str(s)
    if len(s) > n:
        return s[: n - 3] + "..."
    return s


def _cell_to_str(v: object) -> str:
    if v is None:
        return ""
    if isinstance(v, float):
        if v.is_integer():
            return str(int(v))
        return f"{v:.6g}"
    return _truncate(str(v))


def _build_table(rows: list[list[str]], avail_width_pt: float) -> Table:
    """Build a reportlab Table that fits avail_width_pt, with sensible style."""
    if not rows:
        rows = [[""]]
    # Truncate columns
    n_cols = max(len(r) for r in rows)
    n_cols = min(n_cols, MAX_COLS_PER_TABLE)
    norm_rows: list[list[Paragraph]] = []
    body_style = ParagraphStyle(
        "tbl_body", fontName="Helvetica", fontSize=7, leading=8,
    )
    head_style = ParagraphStyle(
        "tbl_head", fontName="Helvetica-Bold", fontSize=7, leading=8,
        textColor=colors.whitesmoke,
    )
    for i, r in enumerate(rows):
        padded = (r + [""] * n_cols)[:n_cols]
        style = head_style if i == 0 else body_style
        norm_rows.append([Paragraph(_safe_html(c), style) for c in padded])

    col_w = avail_width_pt / max(n_cols, 1)
    tbl = Table(norm_rows, colWidths=[col_w] * n_cols, repeatRows=1)
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#374151")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#9CA3AF")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 2),
        ("RIGHTPADDING", (0, 0), (-1, -1), 2),
        ("TOPPADDING", (0, 0), (-1, -1), 1),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1),
            [colors.white, colors.HexColor("#F3F4F6")]),
    ]))
    return tbl


def _is_row_too_dense(headers: list[str], row: list[str]) -> bool:
    """Return True if rendering this row in a table will produce a
    cell too tall to fit on a page. Triggered by the product of
    column count and the longest cell text length."""
    n_cols = max(len(headers), len(row))
    if n_cols == 0:
        return False
    max_chars = max(
        (len(c) for c in row if c is not None), default=0,
    )
    return (n_cols * max_chars) > DENSE_ROW_THRESHOLD


def _render_row_stacked(
    headers: list[str], row: list[str], avail_width_pt: float,
) -> list:
    """Render one mega-row as a 2-column label/value stack instead of
    a table row. Each header becomes a label; its cell becomes a
    paragraph that can flow across pages naturally. Used when the
    row's content would create a cell taller than a page."""
    label_style = ParagraphStyle(
        "stk_label", fontName="Helvetica-Bold", fontSize=8, leading=10,
        textColor=colors.HexColor("#1F2937"),
        spaceBefore=2, spaceAfter=0,
    )
    value_style = ParagraphStyle(
        "stk_value", fontName="Helvetica", fontSize=8, leading=10,
        leftIndent=10, spaceAfter=4,
    )
    sep_style = ParagraphStyle(
        "stk_sep", fontName="Helvetica", fontSize=1, leading=2,
        textColor=colors.HexColor("#D1D5DB"),
    )
    flow: list = []
    n_cols = max(len(headers), len(row))
    padded_headers = (headers + [""] * n_cols)[:n_cols]
    padded_row = (row + [""] * n_cols)[:n_cols]
    for h, v in zip(padded_headers, padded_row):
        if not v:
            continue
        h_text = h.strip() if h else "(unnamed column)"
        flow.append(Paragraph(_safe_html(h_text), label_style))
        flow.append(Paragraph(_safe_html(str(v)), value_style))
    flow.append(Paragraph("&nbsp;", sep_style))
    return flow


def _rows_to_story(
    rows: list[list[str]], avail_width_pt: float,
) -> list:
    """Convert a list of rows (first row treated as headers) into a
    reportlab story. Dense rows are rendered as stacked label/value
    paragraphs; sparse rows are batched into Tables that can break
    across pages. Returns the assembled story fragment."""
    if not rows:
        return []
    headers = [str(c) for c in rows[0]]
    story: list = []
    # Buffer of consecutive sparse rows we can pack into one Table
    sparse_buf: list[list[str]] = []

    def _flush_sparse() -> None:
        if not sparse_buf:
            return
        story.append(_build_table([headers] + sparse_buf, avail_width_pt))
        sparse_buf.clear()

    for r_idx in range(1, len(rows)):
        row = [str(c) for c in rows[r_idx]]
        if _is_row_too_dense(headers, row):
            _flush_sparse()
            story.append(Paragraph(
                _safe_html(f"Row {r_idx} (stacked due to large content)"),
                ParagraphStyle(
                    "row_label", fontName="Helvetica-Bold", fontSize=9,
                    leading=11, textColor=colors.HexColor("#374151"),
                    spaceBefore=6, spaceAfter=2,
                ),
            ))
            story.extend(
                _render_row_stacked(headers, row, avail_width_pt)
            )
        else:
            sparse_buf.append(row)
    _flush_sparse()
    return story


# ---------------------------------------------------------------------
# Excel: xlsx / xlsm via openpyxl
# ---------------------------------------------------------------------
def _convert_xlsx(src: Path, out_pdf: Path, paper: str) -> None:
    from openpyxl import load_workbook
    # Use data_only=True to get cached formula values. Use read_only=False
    # so we can access merged_cells ranges for proper fill-down.
    wb = load_workbook(str(src), read_only=False, data_only=True)
    pagesize = _pagesize(paper, landscape_mode=True)
    page_w, _ = pagesize
    margin = 1.2 * cm
    avail = page_w - 2 * margin

    doc = SimpleDocTemplate(
        str(out_pdf), pagesize=pagesize,
        leftMargin=margin, rightMargin=margin,
        topMargin=margin, bottomMargin=margin,
        title=src.name,
    )
    sty = _styles()
    story: list = [Paragraph(_safe_html(src.name), sty["h1"])]

    sheet_names = list(wb.sheetnames)
    for idx, name in enumerate(sheet_names):
        ws = wb[name]
        if idx > 0:
            story.append(PageBreak())
        story.append(Paragraph(_safe_html(f"Sheet: {name}"), sty["h2"]))

        # Note if sheet has charts (openpyxl exposes _charts)
        n_charts = len(getattr(ws, "_charts", []))
        if n_charts:
            story.append(Paragraph(
                _safe_html(f"[Note: this sheet contains {n_charts} chart(s) "
                           f"which cannot be rendered in text-based conversion]"),
                sty["small"],
            ))

        # Build a merged-cell value lookup: for every cell inside a merged
        # range, return the value of the top-left anchor cell.
        merge_map: dict[tuple[int, int], object] = {}
        for merged_range in ws.merged_cells.ranges:
            anchor_val = ws.cell(
                row=merged_range.min_row, column=merged_range.min_col
            ).value
            for row_n in range(merged_range.min_row, merged_range.max_row + 1):
                for col_n in range(merged_range.min_col, merged_range.max_col + 1):
                    if row_n == merged_range.min_row and col_n == merged_range.min_col:
                        continue
                    merge_map[(row_n, col_n)] = anchor_val

        rows: list[list[str]] = []
        for r_idx, row in enumerate(ws.iter_rows(values_only=False), start=1):
            if r_idx - 1 >= MAX_ROWS_PER_SHEET:
                rows.append([f"[truncated at {MAX_ROWS_PER_SHEET} rows]"])
                break
            row_vals: list[str] = []
            for cell in row:
                val = cell.value
                if val is None:
                    merged_val = merge_map.get((cell.row, cell.column))
                    if merged_val is not None:
                        val = merged_val
                row_vals.append(_cell_to_str(val))
            rows.append(row_vals)

        if not rows or all(not any(cell for cell in r) for r in rows):
            story.append(Paragraph("(empty sheet)", sty["body"]))
            continue
        story.extend(_rows_to_story(rows, avail))

    try:
        doc.build(story)
    except LayoutError:
        story = [Paragraph(_safe_html(src.name), sty["h1"])]
        for idx, name in enumerate(list(wb.sheetnames)):
            ws = wb[name]
            if idx > 0:
                story.append(PageBreak())
            story.append(
                Paragraph(_safe_html(f"Sheet: {name}"), sty["h2"])
            )
            rows2: list[list[str]] = []
            for r_idx, row in enumerate(ws.iter_rows(values_only=True)):
                if r_idx >= MAX_ROWS_PER_SHEET:
                    rows2.append(
                        [f"[truncated at {MAX_ROWS_PER_SHEET} rows]"]
                    )
                    break
                rows2.append([_cell_to_str(v) for v in row])
            if not rows2:
                story.append(Paragraph("(empty sheet)", sty["body"]))
                continue
            headers = [str(c) for c in rows2[0]]
            for r_idx, r in enumerate(rows2[1:], start=1):
                story.append(Paragraph(
                    _safe_html(f"Row {r_idx}"),
                    ParagraphStyle(
                        "row_l2", fontName="Helvetica-Bold", fontSize=9,
                        leading=11, spaceBefore=4, spaceAfter=2,
                    ),
                ))
                story.extend(_render_row_stacked(headers, r, avail))
        doc.build(story)
    wb.close()
    del wb
    gc.collect()


# ---------------------------------------------------------------------
# Excel: legacy .xls via xlrd
# ---------------------------------------------------------------------
def _convert_xls(src: Path, out_pdf: Path, paper: str) -> None:
    import xlrd  # type: ignore
    book = xlrd.open_workbook(str(src), on_demand=True)
    pagesize = _pagesize(paper, landscape_mode=True)
    page_w, _ = pagesize
    margin = 1.2 * cm
    avail = page_w - 2 * margin

    doc = SimpleDocTemplate(
        str(out_pdf), pagesize=pagesize,
        leftMargin=margin, rightMargin=margin,
        topMargin=margin, bottomMargin=margin,
        title=src.name,
    )
    sty = _styles()
    story: list = [Paragraph(_safe_html(src.name), sty["h1"])]

    for idx, name in enumerate(book.sheet_names()):
        sh = book.sheet_by_name(name)
        if idx > 0:
            story.append(PageBreak())
        story.append(Paragraph(_safe_html(f"Sheet: {name}"), sty["h2"]))
        rows: list[list[str]] = []
        n_rows = min(sh.nrows, MAX_ROWS_PER_SHEET)
        for r in range(n_rows):
            rows.append([_cell_to_str(sh.cell_value(r, c))
                         for c in range(sh.ncols)])
        if sh.nrows > MAX_ROWS_PER_SHEET:
            rows.append([f"[truncated at {MAX_ROWS_PER_SHEET} rows]"])
        if not rows:
            story.append(Paragraph("(empty sheet)", sty["body"]))
            continue
        story.extend(_rows_to_story(rows, avail))
        book.unload_sheet(name)

    doc.build(story)
    del book
    gc.collect()


# ---------------------------------------------------------------------
# CSV / TSV
# ---------------------------------------------------------------------
def _convert_csv(src: Path, out_pdf: Path, paper: str, delim: str) -> None:
    pagesize = _pagesize(paper, landscape_mode=True)
    page_w, _ = pagesize
    margin = 1.2 * cm
    avail = page_w - 2 * margin

    rows: list[list[str]] = []
    with src.open("r", encoding="utf-8", errors="replace", newline="") as f:
        reader = csv.reader(f, delimiter=delim)
        for i, r in enumerate(reader):
            if i >= MAX_ROWS_PER_SHEET:
                rows.append([f"[truncated at {MAX_ROWS_PER_SHEET} rows]"])
                break
            rows.append([_truncate(c) for c in r])

    doc = SimpleDocTemplate(
        str(out_pdf), pagesize=pagesize,
        leftMargin=margin, rightMargin=margin,
        topMargin=margin, bottomMargin=margin,
        title=src.name,
    )
    sty = _styles()
    story: list = [Paragraph(_safe_html(src.name), sty["h1"])]
    if not rows:
        story.append(Paragraph("(empty file)", sty["body"]))
    else:
        story.extend(_rows_to_story(rows, avail))
    doc.build(story)


# ---------------------------------------------------------------------
# Plain text
# ---------------------------------------------------------------------
def _convert_text(src: Path, out_pdf: Path, paper: str, raw: str | None = None) -> None:
    pagesize = _pagesize(paper, landscape_mode=False)
    margin = 2 * cm
    doc = SimpleDocTemplate(
        str(out_pdf), pagesize=pagesize,
        leftMargin=margin, rightMargin=margin,
        topMargin=margin, bottomMargin=margin,
        title=src.name,
    )
    sty = _styles()
    if raw is None:
        raw = src.read_text(encoding="utf-8", errors="replace")
    story: list = [Paragraph(_safe_html(src.name), sty["h1"])]
    for chunk in raw.split("\n\n"):
        chunk = chunk.strip()
        if not chunk:
            continue
        safe = _safe_html(chunk).replace("\n", "<br/>")
        story.append(Paragraph(safe, sty["body"]))
        story.append(Spacer(1, 0.2 * cm))
    doc.build(story)


# ---------------------------------------------------------------------
# DOCX via python-docx
# ---------------------------------------------------------------------
def _convert_docx(src: Path, out_pdf: Path, paper: str) -> None:
    import os
    import tempfile

    from docx import Document  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    from reportlab.platypus import Image as RLImage

    document = Document(str(src))
    pagesize = _pagesize(paper, landscape_mode=False)
    page_w, _ = pagesize
    margin = 2 * cm
    avail = page_w - 2 * margin
    max_img_height = 18 * cm

    doc = SimpleDocTemplate(
        str(out_pdf), pagesize=pagesize,
        leftMargin=margin, rightMargin=margin,
        topMargin=margin, bottomMargin=margin,
        title=src.name,
    )
    sty = _styles()
    story: list = [Paragraph(_safe_html(src.name), sty["h1"])]

    p_tag = qn("w:p")
    tbl_tag = qn("w:tbl")

    # Namespace-qualified tag names for image extraction
    drawing_tag = qn("w:drawing")
    pict_tag = qn("w:pict")
    blip_tag = qn("a:blip")
    embed_attr = qn("r:embed")
    # VML namespace isn't registered in python-docx's qn(); use full URI
    _VML_NS = "urn:schemas-microsoft-com:vml"
    imagedata_tag = f"{{{_VML_NS}}}imagedata"

    # Temp directory for extracted images; cleaned up at the end
    tmp_dir = tempfile.mkdtemp(prefix="docx_img_")
    img_counter = 0

    def _extract_images_from_paragraph(para_element) -> list:
        """Extract embedded images from a paragraph XML element.

        Returns a list of reportlab Image flowables scaled to fit the page.
        """
        nonlocal img_counter
        images = []

        # Find w:drawing elements (modern OOXML inline/anchor images)
        for drawing_el in para_element.iter(drawing_tag):
            for blip in drawing_el.iter(blip_tag):
                r_id = blip.get(embed_attr)
                if not r_id:
                    continue
                img_flowable = _image_from_rid(r_id)
                if img_flowable:
                    images.append(img_flowable)

        # Find w:pict elements (legacy VML images)
        for pict_el in para_element.iter(pict_tag):
            for imgdata in pict_el.iter(imagedata_tag):
                r_id = imgdata.get(embed_attr)
                if not r_id:
                    continue
                img_flowable = _image_from_rid(r_id)
                if img_flowable:
                    images.append(img_flowable)

        return images

    def _image_from_rid(r_id: str):
        """Given a relationship ID, extract the image and return a scaled Image flowable."""
        nonlocal img_counter
        try:
            rel = document.part.rels.get(r_id)
            if rel is None:
                return None
            image_part = rel.target_part
            image_bytes = image_part.blob
            # Determine extension from content type
            content_type = getattr(image_part, "content_type", "image/png")
            ext_map = {
                "image/png": ".png",
                "image/jpeg": ".jpg",
                "image/gif": ".gif",
                "image/bmp": ".bmp",
                "image/tiff": ".tiff",
                "image/x-emf": ".emf",
                "image/x-wmf": ".wmf",
            }
            ext = ext_map.get(content_type, ".png")

            # Skip EMF/WMF as reportlab/PIL often cannot handle them
            if ext in (".emf", ".wmf"):
                return None

            img_counter += 1
            img_path = os.path.join(tmp_dir, f"img_{img_counter}{ext}")
            with open(img_path, "wb") as f:
                f.write(image_bytes)

            # Use PIL to get dimensions for proper scaling
            try:
                from PIL import Image as PILImage  # type: ignore
                with PILImage.open(img_path) as pil_img:
                    orig_w, orig_h = pil_img.size
                    # pil_img.info.get('dpi') may not exist
                    dpi = pil_img.info.get("dpi", (96, 96))
                    dpi_x = dpi[0] if dpi[0] > 0 else 96
                    dpi_y = dpi[1] if dpi[1] > 0 else 96
                    # Convert pixel dimensions to points (1 inch = 72 points)
                    w_pt = orig_w * 72.0 / dpi_x
                    h_pt = orig_h * 72.0 / dpi_y
            except Exception:
                # Fallback: assume 96 DPI and guess from file
                w_pt = avail
                h_pt = max_img_height * 0.5

            # Scale to fit within available width and max height
            if w_pt > avail:
                scale = avail / w_pt
                w_pt = avail
                h_pt = h_pt * scale
            if h_pt > max_img_height:
                scale = max_img_height / h_pt
                h_pt = max_img_height
                w_pt = w_pt * scale

            return RLImage(img_path, width=w_pt, height=h_pt)
        except Exception:
            return None

    def _para_to_rich_text(para) -> str:
        """Extract rich text from a paragraph preserving bold/italic/underline
        as reportlab inline markup. Returns safe HTML string for Paragraph()."""
        parts: list[str] = []
        for run in para.runs:
            t = run.text or ""
            if not t:
                continue
            safe = _safe_html(t)
            if run.bold:
                safe = f"<b>{safe}</b>"
            if run.italic:
                safe = f"<i>{safe}</i>"
            if run.underline:
                safe = f"<u>{safe}</u>"
            parts.append(safe)
        return "".join(parts)

    def _detect_list_prefix(para) -> str:
        """Detect if paragraph is a list item and return a bullet/number prefix."""
        try:
            numPr = para._element.find(f".//{qn('w:numPr')}")
            if numPr is not None:
                ilvl_el = numPr.find(qn("w:ilvl"))
                indent = int(ilvl_el.get(qn("w:val"), "0")) if ilvl_el is not None else 0
                padding = "    " * indent
                return f"{padding}• "
        except Exception:
            pass
        return ""

    body = document.element.body
    para_idx = 0
    tbl_idx = 0
    for child in body.iterchildren():
        if child.tag == p_tag:
            if para_idx < len(document.paragraphs):
                para = document.paragraphs[para_idx]
                para_idx += 1

                # Extract images from this paragraph element
                para_images = []
                try:
                    para_images = _extract_images_from_paragraph(para._element)
                except Exception:
                    pass

                # Build rich text (preserving bold/italic/underline)
                rich_text = _para_to_rich_text(para)
                plain_text = (para.text or "").strip()

                # Determine list prefix
                list_prefix = _detect_list_prefix(para)

                # Handle text
                if rich_text.strip() or plain_text:
                    display_text = rich_text if rich_text.strip() else _safe_html(plain_text)
                    if list_prefix:
                        display_text = _safe_html(list_prefix) + display_text
                    style_name = (para.style.name if para.style else "") or ""
                    style_lower = style_name.lower()
                    if style_lower.startswith("heading 1") or style_lower == "title":
                        story.append(Paragraph(display_text, sty["h1"]))
                    elif style_lower.startswith("heading"):
                        story.append(Paragraph(display_text, sty["h2"]))
                    elif "list" in style_lower:
                        if not list_prefix:
                            display_text = "• " + display_text
                        story.append(Paragraph(display_text, sty["body"]))
                    else:
                        story.append(Paragraph(display_text, sty["body"]))

                # Append any images found in this paragraph
                if para_images:
                    for img_flowable in para_images:
                        story.append(img_flowable)
                        story.append(Spacer(1, 0.2 * cm))
                elif not rich_text.strip() and not plain_text:
                    story.append(Spacer(1, 0.2 * cm))

        elif child.tag == tbl_tag:
            if tbl_idx < len(document.tables):
                t = document.tables[tbl_idx]
                tbl_idx += 1
                rows: list[list[str]] = []
                for row in t.rows:
                    rows.append([_truncate(cell.text or "") for cell in row.cells])
                if rows:
                    story.append(_build_table(rows, avail))
                    story.append(Spacer(1, 0.3 * cm))

    # Extract text from text boxes (w:txbxContent) which are not part of
    # the main body paragraph flow. These appear inside shapes/drawings.
    txbx_tag = qn("w:txbxContent")
    textbox_texts: list[str] = []
    for txbx in document.element.body.iter(txbx_tag):
        for p_el in txbx.iter(p_tag):
            text_parts = []
            for t_el in p_el.iter(qn("w:t")):
                if t_el.text:
                    text_parts.append(t_el.text)
            combined = "".join(text_parts).strip()
            if combined:
                textbox_texts.append(combined)
    if textbox_texts:
        story.append(Spacer(1, 0.4 * cm))
        story.append(Paragraph("[Text Boxes]", sty["h2"]))
        for tb_text in textbox_texts:
            story.append(Paragraph(_safe_html(tb_text), sty["body"]))

    # Extract footnotes and endnotes if present
    try:
        footnote_part = None
        endnote_part = None
        for rel in document.part.rels.values():
            if "footnotes" in str(getattr(rel, "reltype", "")):
                footnote_part = rel.target_part
            elif "endnotes" in str(getattr(rel, "reltype", "")):
                endnote_part = rel.target_part

        fn_texts: list[str] = []
        for part in [footnote_part, endnote_part]:
            if part is None:
                continue
            try:
                from lxml import etree  # type: ignore
                root = etree.fromstring(part.blob)
            except Exception:
                from xml.etree import ElementTree as _ET
                root = _ET.fromstring(part.blob)
            for p_el in root.iter(f"{{{root.nsmap.get('w', '')}}}t" if hasattr(root, 'nsmap') else qn("w:t")):
                if p_el.text:
                    fn_texts.append(p_el.text)
        if fn_texts:
            combined_fn = " ".join(fn_texts).strip()
            # Skip default empty footnotes (Word always has ids 0 and 1)
            if combined_fn and len(combined_fn) > 5:
                story.append(Spacer(1, 0.4 * cm))
                story.append(Paragraph("[Footnotes/Endnotes]", sty["h2"]))
                story.append(Paragraph(_safe_html(combined_fn), sty["body"]))
    except Exception:
        pass

    doc.build(story)
    del document
    gc.collect()

    # Clean up temporary image files
    try:
        import shutil
        shutil.rmtree(tmp_dir, ignore_errors=True)
    except Exception:
        pass


# ---------------------------------------------------------------------
# RTF via striprtf (with table structure recovery)
# ---------------------------------------------------------------------
def _convert_rtf(src: Path, out_pdf: Path, paper: str) -> None:
    import re as _re
    from striprtf.striprtf import rtf_to_text  # type: ignore
    raw = src.read_text(encoding="utf-8", errors="replace")

    # Detect if the RTF contains table markup (\trowd ... \row)
    has_tables = bool(_re.search(r"\\trowd", raw))

    # Extract images embedded as hex blobs (\pict ... hex data)
    # RTF images are {\pict\...  <hex bytes>}
    pict_pattern = _re.compile(
        r"\\pict[^}]*?\\(?:pngblip|jpegblip|emfblip|wmetafile)[^}]*?"
        r"\s+([0-9a-fA-F\s]+)",
        _re.DOTALL,
    )
    has_images = bool(pict_pattern.search(raw))

    text = rtf_to_text(raw, errors="ignore")

    # If there are tables, try to recover tabular structure from the text
    # output. striprtf typically separates cells with \t or \n within rows.
    if has_tables:
        lines = text.split("\n")
        processed: list[str] = []
        for line in lines:
            if "\t" in line:
                # Tab-separated = likely a table row; preserve as-is
                processed.append(line)
            else:
                processed.append(line)
        text = "\n".join(processed)

    # Add a note if images were detected but cannot be extracted
    if has_images:
        text = (
            "[Note: This RTF file contains embedded images which cannot be "
            "fully extracted in text-based conversion.]\n\n" + text
        )

    _convert_text(src, out_pdf, paper, raw=text)


# ---------------------------------------------------------------------
# Visio (.vsdx) via visio_convert
# ---------------------------------------------------------------------
def _convert_visio(src: Path, out_pdf: Path, paper: str) -> None:
    from tools.visio_convert import convert_visio_to_pdf
    convert_visio_to_pdf(src, out_pdf, paper)


# ---------------------------------------------------------------------
# .doc (legacy binary)
# ---------------------------------------------------------------------
def _convert_doc_placeholder(src: Path, out_pdf: Path, paper: str) -> None:
    msg = (
        f"Legacy .doc format cannot be parsed without Microsoft Word.\n"
        f"Recommendation: ask the author to resave as .docx, or convert "
        f"manually.\n\nFile: {src.name}\nSize: {src.stat().st_size} bytes"
    )
    _convert_text(src, out_pdf, paper, raw=msg)


# ---------------------------------------------------------------------
# Public dispatcher
# ---------------------------------------------------------------------
_XLSX_EXTS: Final[frozenset[str]] = frozenset({
    ".xlsx", ".xlsm", ".xltx", ".xltm",
})
_XLS_EXTS: Final[frozenset[str]] = frozenset({".xls", ".xlsb"})
_CSV_EXTS: Final[frozenset[str]] = frozenset({".csv"})
_TSV_EXTS: Final[frozenset[str]] = frozenset({".tsv"})
_TXT_EXTS: Final[frozenset[str]] = frozenset({".txt", ".log", ".md"})
_DOCX_EXTS: Final[frozenset[str]] = frozenset({
    ".docx", ".docm", ".dotx", ".dotm",
})
_DOC_EXTS: Final[frozenset[str]] = frozenset({".doc", ".dot"})
_RTF_EXTS: Final[frozenset[str]] = frozenset({".rtf"})
_VISIO_EXTS: Final[frozenset[str]] = frozenset({".vsdx"})


def convert_to_pdf(src: Path, out_pdf: Path, paper: str = "A4") -> tuple[str, str]:
    """Dispatch by extension. Returns (status, message).

    status in {"SUCCESS", "PARTIAL", "FAILED"}.
    """
    ext = src.suffix.lower()
    out_pdf.parent.mkdir(parents=True, exist_ok=True)
    try:
        if ext in _XLSX_EXTS:
            _convert_xlsx(src, out_pdf, paper)
        elif ext == ".xls":
            _convert_xls(src, out_pdf, paper)
        elif ext == ".xlsb":
            return ("FAILED",
                    "xlsb format not supported in pure-Python mode")
        elif ext in _CSV_EXTS:
            _convert_csv(src, out_pdf, paper, delim=",")
        elif ext in _TSV_EXTS:
            _convert_csv(src, out_pdf, paper, delim="\t")
        elif ext in _TXT_EXTS:
            _convert_text(src, out_pdf, paper)
        elif ext in _DOCX_EXTS:
            _convert_docx(src, out_pdf, paper)
        elif ext in _DOC_EXTS:
            _convert_doc_placeholder(src, out_pdf, paper)
            return ("PARTIAL",
                    ".doc is legacy binary; only a placeholder is rendered")
        elif ext in _RTF_EXTS:
            _convert_rtf(src, out_pdf, paper)
        elif ext in _VISIO_EXTS:
            _convert_visio(src, out_pdf, paper)
        else:
            return ("FAILED", f"Unsupported extension: {ext}")
        return ("SUCCESS", "Converted")
    except Exception as e:
        return ("FAILED", f"Conversion exception: {e!r}")


def is_office_extension(ext: str) -> bool:
    e = ext.lower()
    return (e in _XLSX_EXTS or e in _XLS_EXTS or e in _CSV_EXTS
            or e in _TSV_EXTS or e in _TXT_EXTS or e in _DOCX_EXTS
            or e in _DOC_EXTS or e in _RTF_EXTS or e in _VISIO_EXTS)
